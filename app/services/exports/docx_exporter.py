"""Microsoft Word exports for Filtrify campaign assets."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from app.services.exports.base_exporter import BaseExporter


class DocxExporter(BaseExporter):
    """Create professional DOCX documents for Filtrify campaign assets."""

    def seo_article_to_bytes(self, article: Any) -> bytes:
        """Return a formatted SEO article DOCX."""

        document = self._new_document(
            title=article.title,
            subtitle=f"Target keyword: {article.target_keyword}",
        )

        document.add_heading("Meta Description", level=1)
        document.add_paragraph(article.meta_description)

        document.add_heading("Introduction", level=1)
        document.add_paragraph(article.introduction)

        for section in article.sections:
            document.add_heading(section.heading, level=1)
            self._add_text_with_bullets(
                document,
                section.content,
            )

        document.add_heading("Conclusion", level=1)
        document.add_paragraph(article.conclusion)

        document.add_heading("Call to Action", level=1)
        document.add_paragraph(article.call_to_action)

        document.add_section()
        document.add_paragraph(
            f"SEO score: {article.seo_score:.1f}/100"
        )
        document.add_paragraph(
            f"Estimated words: {article.estimated_word_count:,}"
        )

        return self._document_to_bytes(document)

    def landing_page_to_bytes(self, landing_page: Any) -> bytes:
        """Return a formatted landing-page DOCX."""

        document = self._new_document(
            title=landing_page.page_title,
            subtitle=(
                f"Audience: {landing_page.target_audience} | "
                f"Goal: {landing_page.primary_goal}"
            ),
        )

        document.add_heading("Meta Description", level=1)
        document.add_paragraph(landing_page.meta_description)

        document.add_heading("Hero Section", level=1)
        document.add_heading(landing_page.hero_headline, level=2)
        document.add_paragraph(landing_page.hero_subheadline)
        document.add_paragraph(
            f"Primary CTA: {landing_page.primary_cta}"
        )

        for section in landing_page.sections:
            document.add_heading(section.heading, level=1)
            document.add_paragraph(section.content)

            for item in section.items:
                document.add_paragraph(
                    item,
                    style="List Bullet",
                )

        document.add_heading(
            landing_page.final_cta_heading,
            level=1,
        )
        document.add_paragraph(landing_page.final_cta_text)
        document.add_paragraph(
            f"Final CTA: {landing_page.final_cta_button}"
        )

        document.add_paragraph(
            f"Conversion score: "
            f"{landing_page.conversion_score:.1f}/100"
        )
        document.add_paragraph(
            f"Estimated words: "
            f"{landing_page.estimated_word_count:,}"
        )

        return self._document_to_bytes(document)

    def email_sequence_to_bytes(self, email_sequence: Any) -> bytes:
        """Return a formatted email-sequence DOCX."""

        document = self._new_document(
            title=email_sequence.sequence_name,
            subtitle=(
                f"Product: {email_sequence.product_name} | "
                f"Audience: {email_sequence.target_audience}"
            ),
        )

        document.add_heading("Strategy Summary", level=1)
        document.add_paragraph(email_sequence.strategy_summary)
        document.add_paragraph(
            f"Primary goal: {email_sequence.primary_goal}"
        )

        for email in email_sequence.emails:
            document.add_page_break()
            document.add_heading(
                f"Email {email.sequence_number}: {email.purpose}",
                level=1,
            )

            self._add_labelled_paragraph(
                document,
                "Subject",
                email.subject,
            )
            self._add_labelled_paragraph(
                document,
                "Preview text",
                email.preview_text,
            )

            document.add_heading("Email Body", level=2)
            self._add_text_with_bullets(
                document,
                email.body,
            )

            self._add_labelled_paragraph(
                document,
                "Call to Action",
                email.call_to_action,
            )

        return self._document_to_bytes(document)

    def campaign_summary_to_bytes(self, campaign: Any) -> bytes:
        """Return a summary DOCX for a complete campaign."""

        self.validate_campaign(campaign)

        document = self._new_document(
            title=campaign.campaign_name,
            subtitle=f"Product: {campaign.product_name}",
        )

        document.add_heading("Campaign Summary", level=1)

        summary_items = (
            ("Assets", str(campaign.asset_count)),
            (
                "Estimated words",
                f"{campaign.total_estimated_words:,}",
            ),
            (
                "Average quality",
                f"{campaign.average_quality_score:.1f}/100",
            ),
            ("Target keyword", campaign.target_keyword),
            ("Target audience", campaign.target_audience),
            ("Tone", campaign.tone),
            (
                "Created",
                campaign.created_at.isoformat(),
            ),
        )

        table = document.add_table(
            rows=1,
            cols=2,
        )
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        header_cells[0].text = "Field"
        header_cells[1].text = "Value"

        for label, value in summary_items:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value

        document.add_heading("Included Assets", level=1)
        for item in (
            "SEO Article",
            "Landing Page",
            "Email Sequence",
            "Google Ads Campaign",
        ):
            document.add_paragraph(
                item,
                style="List Bullet",
            )

        return self._document_to_bytes(document)

    def export_seo_article(
        self,
        article: Any,
        *,
        output_directory: str | Path = "outputs/campaigns",
        file_name: str | None = None,
    ) -> Path:
        """Write an SEO article DOCX to disk."""

        resolved_name = file_name or self.build_file_name(
            stem=f"{article.title}_seo_article",
            extension="docx",
        )

        return self.write_bytes(
            data=self.seo_article_to_bytes(article),
            output_directory=output_directory,
            file_name=self._ensure_docx_name(resolved_name),
        )

    def export_landing_page(
        self,
        landing_page: Any,
        *,
        output_directory: str | Path = "outputs/campaigns",
        file_name: str | None = None,
    ) -> Path:
        """Write a landing-page DOCX to disk."""

        resolved_name = file_name or self.build_file_name(
            stem=f"{landing_page.product_name}_landing_page",
            extension="docx",
        )

        return self.write_bytes(
            data=self.landing_page_to_bytes(landing_page),
            output_directory=output_directory,
            file_name=self._ensure_docx_name(resolved_name),
        )

    def export_email_sequence(
        self,
        email_sequence: Any,
        *,
        output_directory: str | Path = "outputs/campaigns",
        file_name: str | None = None,
    ) -> Path:
        """Write an email-sequence DOCX to disk."""

        resolved_name = file_name or self.build_file_name(
            stem=email_sequence.sequence_name,
            extension="docx",
        )

        return self.write_bytes(
            data=self.email_sequence_to_bytes(email_sequence),
            output_directory=output_directory,
            file_name=self._ensure_docx_name(resolved_name),
        )

    @staticmethod
    def _new_document(
        *,
        title: str,
        subtitle: str = "",
    ) -> Any:
        """Create a document with shared Filtrify formatting."""

        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches, Pt
        except ImportError as error:
            raise RuntimeError(
                "python-docx is required for DOCX exports. "
                "Install it with: pip install python-docx"
            ) from error

        document = Document()

        section = document.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

        styles = document.styles
        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(10.5)

        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(20)

        if subtitle:
            subtitle_paragraph = document.add_paragraph()
            subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_paragraph.add_run(subtitle)
            subtitle_run.italic = True
            subtitle_run.font.size = Pt(9)

        document.add_paragraph()
        return document

    @staticmethod
    def _add_labelled_paragraph(
        document: Any,
        label: str,
        value: str,
    ) -> None:
        """Add a bold label followed by normal text."""

        paragraph = document.add_paragraph()
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        paragraph.add_run(value)

    @staticmethod
    def _add_text_with_bullets(
        document: Any,
        content: str,
    ) -> None:
        """Add paragraphs while preserving simple bullet lines."""

        for raw_block in str(content).split("\n"):
            block = raw_block.strip()

            if not block:
                continue

            if block.startswith(("- ", "• ")):
                document.add_paragraph(
                    block[2:].strip(),
                    style="List Bullet",
                )
            else:
                document.add_paragraph(block)

    @staticmethod
    def _document_to_bytes(document: Any) -> bytes:
        """Serialize a python-docx document into bytes."""

        stream = BytesIO()
        document.save(stream)
        stream.seek(0)
        return stream.read()

    @classmethod
    def _ensure_docx_name(cls, file_name: str) -> str:
        """Ensure a safe DOCX file name."""

        if file_name.lower().endswith(".docx"):
            stem = file_name[:-5]
        else:
            stem = file_name

        return cls.build_file_name(
            stem=stem,
            extension="docx",
        )

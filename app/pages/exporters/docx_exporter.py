from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.exporters.content_exporter import (
    ExportedContent,
)


class DOCXExporter:
    """
    Export generated content as a Microsoft Word document.

    DOCX = Microsoft Word Open XML Document.

    XML = Extensible Markup Language.

    DOCX files can be opened and edited using applications
    such as Microsoft Word, Google Docs and LibreOffice.
    """

    @property
    def format_key(self) -> str:
        """
        Return the internal export-format identifier.
        """
        return "docx"

    @property
    def display_name(self) -> str:
        """
        Return the format name shown to users.
        """
        return "DOCX"

    def export(
        self,
        *,
        title: str,
        content: str,
        filename_stem: str,
    ) -> ExportedContent:
        """
        Convert generated content into a Word document.
        """
        document = Document()

        self._configure_document_styles(
            document
        )

        cleaned_title = title.strip()

        if cleaned_title:
            title_paragraph = document.add_paragraph()

            title_paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            title_run = title_paragraph.add_run(
                cleaned_title
            )

            title_run.bold = True
            title_run.font.size = Pt(20)

            document.add_paragraph()

        self._add_content(
            document=document,
            content=content,
        )

        output_buffer = BytesIO()

        document.save(
            output_buffer
        )

        output_buffer.seek(0)

        return ExportedContent(
            filename=f"{filename_stem}.docx",
            data=output_buffer.getvalue(),
            mime_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )

    @staticmethod
    def _configure_document_styles(
        document: Document,
    ) -> None:
        """
        Configure the default Word document styling.
        """
        normal_style = document.styles[
            "Normal"
        ]

        normal_style.font.name = "Arial"
        normal_style.font.size = Pt(11)

    @staticmethod
    def _add_content(
        *,
        document: Document,
        content: str,
    ) -> None:
        """
        Add plain text to the Word document.

        Blank lines create spacing.

        Lines beginning with "- " become bullet points.
        """
        for raw_line in content.splitlines():
            cleaned_line = raw_line.strip()

            if not cleaned_line:
                document.add_paragraph()
                continue

            if cleaned_line.startswith("- "):
                bullet_text = cleaned_line[2:].strip()

                document.add_paragraph(
                    bullet_text,
                    style="List Bullet",
                )

                continue

            document.add_paragraph(
                cleaned_line
            )